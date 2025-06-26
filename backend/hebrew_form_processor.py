#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Hebrew Physiotherapy Workshops Form Generator
Processes Hebrew text to extract participant information and generate Word documents
"""

import re
import sys
import logging
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional, Tuple

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from groq_client import GroqClient

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class HebrewFormProcessor:
    """Process Hebrew text and generate Word documents for physiotherapy workshops."""
    
    # Activity type mapping
    ACTIVITY_TYPES = {
        'התעמלות לאחר לידה': 'post_birth_exercise',
        'התעמלות הריון': 'pregnancy_exercise'
    }
    
    # Hebrew month names
    HEBREW_MONTHS = {
        1: 'ינואר', 2: 'פברואר', 3: 'מרץ', 4: 'אפריל',
        5: 'מאי', 6: 'יוני', 7: 'יולי', 8: 'אוגוסט',
        9: 'ספטמבר', 10: 'אוקטובר', 11: 'נובמבר', 12: 'דצמבר'
    }
    
    def __init__(self, groq_api_key: Optional[str] = None):
        """Initialize the Hebrew Form Processor."""
        self.activity_type = None
        self.participants = []
        self.groq_client = GroqClient(groq_api_key)
    
    def identify_activity_type(self, text: str) -> Optional[str]:
        """
        Identify the activity type from the text title.
        
        Args:
            text: The input Hebrew text
            
        Returns:
            The activity type or None if not found
        """
        # First try with Groq API for intelligent identification
        try:
            groq_result = self.groq_client.identify_activity_type(text)
            if groq_result:
                logger.info(f"Activity type identified by Groq: {groq_result}")
                return groq_result
        except Exception as e:
            logger.warning(f"Groq API failed for activity identification: {str(e)}")
        
        # Fallback to regex pattern matching
        for hebrew_activity, activity_code in self.ACTIVITY_TYPES.items():
            if hebrew_activity in text:
                logger.info(f"Activity type identified by regex: {hebrew_activity}")
                return hebrew_activity
        
        logger.warning("No activity type identified")
        return None
    
    def extract_participant_info(self, text: str) -> List[Dict[str, str]]:
        """
        Extract participant information from Hebrew text.
        
        Args:
            text: The input Hebrew text
            
        Returns:
            List of dictionaries containing participant information
        """
        participants = []
        
        # First try with Groq API for intelligent extraction
        try:
            groq_participants = self.groq_client.extract_participant_info(text)
            if groq_participants:
                logger.info(f"Groq API extracted {len(groq_participants)} participants")
                return groq_participants
        except Exception as e:
            logger.warning(f"Groq API failed for participant extraction: {str(e)}")
        
        # Fallback to regex-based extraction for structured text
        logger.info("Falling back to regex-based extraction")
        return self._extract_with_regex(text)
    
    def _extract_with_regex(self, text: str) -> List[Dict[str, str]]:
        """
        Extract participant information using regex patterns (fallback method).
        
        Args:
            text: The input Hebrew text
            
        Returns:
            List of dictionaries containing participant information
        """
        participants = []
        
        # Pattern to match Hebrew names (Hebrew letters, spaces, and common punctuation)
        hebrew_name_pattern = r'[\u0590-\u05FF\s\-\'\"]{2,50}'
        
        # Pattern to match Israeli ID numbers (9 digits)
        id_pattern = r'\b\d{9}\b'
        
        # Pattern to match receipt numbers (various formats)
        receipt_pattern = r'(?:קבלה|קבלה מס\'|מס\' קבלה|receipt|חשבונית)[:\s]*(\d+)'
        
        # Pattern to match amounts (numbers with currency symbols or Hebrew currency words)
        amount_pattern = r'(?:₪|שקל|שקלים|ש״ח)[:\s]*(\d+(?:\.\d{2})?)|(\d+(?:\.\d{2})?)[:\s]*(?:₪|שקל|שקלים|ש״ח)'
        
        # Enhanced patterns for natural language extraction
        name_context_patterns = [
            r'(?:שם|משתתף|משתתפת|השתתף|השתתפה)[:\s]*([\u0590-\u05FF\s\-\'\"]{2,30})',
            r'([\u0590-\u05FF]+\s+[\u0590-\u05FF]+)(?:\s+השתתף|\s+השתתפה|\s+שילם|\s+שילמה)',
        ]
        
        id_context_patterns = [
            r'(?:תעודת זהות|ת\.ז|ת״ז|זהות)[:\s]*(\d{9})',
            r'(\d{9})(?=.*(?:תעודת זהות|ת\.ז|ת״ז))',
        ]
        
        amount_context_patterns = [
            r'(?:שילם|שילמה|תשלום|סכום)[:\s]*(\d+)(?:\s*(?:₪|שקל|שקלים|ש״ח))?',
            r'(\d+)[:\s]*(?:₪|שקל|שקלים|ש״ח)',
        ]
        
        # Split text into sentences for better processing
        sentences = re.split(r'[.\n]+', text)
        
        current_participant = {}
        
        for sentence in sentences:
            sentence = sentence.strip()
            if not sentence:
                continue
            
            # Try to extract name with context
            if not current_participant.get('name'):
                for pattern in name_context_patterns:
                    name_match = re.search(pattern, sentence, re.IGNORECASE)
                    if name_match:
                        potential_name = name_match.group(1).strip()
                        # Skip if it's an activity type
                        if not any(activity in potential_name for activity in self.ACTIVITY_TYPES.keys()):
                            current_participant['name'] = potential_name
                            break
            
            # Extract ID with context
            if not current_participant.get('id'):
                for pattern in id_context_patterns:
                    id_match = re.search(pattern, sentence, re.IGNORECASE)
                    if id_match:
                        current_participant['id'] = id_match.group(1)
                        break
            
            # Extract receipt number
            receipt_match = re.search(receipt_pattern, sentence, re.IGNORECASE)
            if receipt_match and not current_participant.get('receipt_number'):
                current_participant['receipt_number'] = receipt_match.group(1)
            
            # Extract amount with context
            if not current_participant.get('amount'):
                for pattern in amount_context_patterns:
                    amount_match = re.search(pattern, sentence, re.IGNORECASE)
                    if amount_match:
                        current_participant['amount'] = amount_match.group(1)
                        break
            
            # If we have collected information for a participant, add to list
            if current_participant.get('name') and len(current_participant) >= 2:
                # Check if we have enough information
                participants.append(current_participant.copy())
                current_participant = {}
        
        # Add any remaining participant if they have at least name
        if current_participant.get('name'):
            participants.append(current_participant)
        
        # Fill in missing fields with empty strings
        for participant in participants:
            for field in ['name', 'id', 'receipt_number', 'amount']:
                if field not in participant:
                    participant[field] = ''
        
        logger.info(f"Regex extraction found {len(participants)} participants")
        return participants
    
    def process_text(self, text: str) -> Dict:
        """
        Process the input text to extract all relevant information.
        
        Args:
            text: The input Hebrew text
            
        Returns:
            Dictionary containing processed information
        """
        logger.info("Starting text processing")
        
        # Identify activity type
        self.activity_type = self.identify_activity_type(text)
        
        # Extract participant information
        self.participants = self.extract_participant_info(text)
        
        result = {
            'activity_type': self.activity_type,
            'participants': self.participants,
            'total_participants': len(self.participants)
        }
        
        logger.info(f"Processing complete: {result['total_participants']} participants found")
        return result
    
    def create_word_document(self, output_path: str = None, custom_title: str = None) -> str:
        """
        Create a Word document with the extracted information.
        
        Args:
            output_path: Path where to save the document
            custom_title: Custom title for the document (overrides default)
            
        Returns:
            Path to the created document
        """
        if not self.participants:
            raise ValueError("No participant information found. Please process text first.")
        
        # Create a new document
        doc = Document()
        
        # Set document direction to RTL for Hebrew
        doc.sections[0].start_type = 1
        
        # Create title
        if custom_title:
            title = custom_title
        else:
            # Default title with current month
            current_month = self.HEBREW_MONTHS[datetime.now().month]
            
            # Use the first participant's name for the title, or a generic title
            if self.participants and self.participants[0].get('name'):
                title_name = self.participants[0]['name']
            else:
                title_name = "משתתפים"
            
            title = f"{title_name} - {current_month}"
        
        # Add title to document
        title_paragraph = doc.add_paragraph()
        title_run = title_paragraph.add_run(title)
        title_run.font.size = Inches(0.2)
        title_run.bold = True
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add some spacing
        doc.add_paragraph()
        
        # Create table
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        
        # Set table direction to RTL
        table.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Add header row
        header_cells = table.rows[0].cells
        #headers = ['שם', 'תעודת זהות', 'מספר קבלה', 'סכום ששולם']
        headers = ['סכום', 'מסםר קבלה', 'תעודת זהות', 'שם']
        
        for i, header in enumerate(headers):
            header_cells[i].text = header
            # Make header bold
            for paragraph in header_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add participant data
        for participant in self.participants:
            row_cells = table.add_row().cells
            # row_cells[0].text = participant.get('name', '')
            # row_cells[1].text = participant.get('id', '')
            # row_cells[2].text = participant.get('receipt_number', '')
            # row_cells[3].text = participant.get('amount', '')
            row_cells[0].text = participant.get('amount', '')
            row_cells[1].text = participant.get('receipt_number', '')
            row_cells[2].text = participant.get('id', '')
            row_cells[3].text = participant.get('name', '')
            
            # Center align the data
            for cell in row_cells:
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Generate output filename if not provided
        if not output_path:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            activity_suffix = "post_birth" if "לאחר לידה" in str(self.activity_type) else "pregnancy"
            output_path = f"physiotherapy_form_{activity_suffix}_{timestamp}.docx"
        
        # Save the document
        doc.save(output_path)
        logger.info(f"Word document saved: {output_path}")
        return output_path 