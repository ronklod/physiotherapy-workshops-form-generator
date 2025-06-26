#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Hebrew Physiotherapy Workshops Form Generator
Processes Hebrew text to extract participant information and generate Word documents
"""

import re
import sys
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional, Tuple

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


class HebrewFormProcessor:
    """Process Hebrew text and generate Word documents for physiotherapy workshops."""
    
    # Activity type mapping
    ACTIVITY_TYPES = {
        'התעמלות לאחר לידה': 'post_birth_exercise',
        'התעמלות הריון': 'pregnancy_exercise'
    }
    
    # Hebrew month names
    HEBREW_MONTHS = {
        1: 'ינואר', 2: 'פברואר', 3: 'מרץ', 4: 'אפריל',
        5: 'מאי', 6: 'יוני', 7: 'יולי', 8: 'אוגוסט',
        9: 'ספטמבר', 10: 'אוקטובר', 11: 'נובמבר', 12: 'דצמבר'
    }
    
    def __init__(self):
        """Initialize the Hebrew Form Processor."""
        self.activity_type = None
        self.participants = []
    
    def identify_activity_type(self, text: str) -> Optional[str]:
        """
        Identify the activity type from the text title.
        
        Args:
            text: The input Hebrew text
            
        Returns:
            The activity type or None if not found
        """
        for hebrew_activity, activity_code in self.ACTIVITY_TYPES.items():
            if hebrew_activity in text:
                return hebrew_activity
        return None
    
    def extract_participant_info(self, text: str) -> List[Dict[str, str]]:
        """
        Extract participant information from Hebrew text.
        
        Args:
            text: The input Hebrew text
            
        Returns:
            List of dictionaries containing participant information
        """
        participants = []
        
        # Pattern to match Hebrew names (Hebrew letters, spaces, and common punctuation)
        hebrew_name_pattern = r'[\u0590-\u05FF\s\-\'\"]{2,50}'
        
        # Pattern to match Israeli ID numbers (9 digits)
        id_pattern = r'\b\d{9}\b'
        
        # Pattern to match receipt numbers (various formats)
        receipt_pattern = r'(?:קבלה|קבלה מס\'|מס\' קבלה|receipt)[:\s]*(\d+)'
        
        # Pattern to match amounts (numbers with currency symbols or Hebrew currency words)
        amount_pattern = r'(?:₪|שקל|שקלים|ש״ח)[:\s]*(\d+(?:\.\d{2})?)|(\d+(?:\.\d{2})?)[:\s]*(?:₪|שקל|שקלים|ש״ח)'
        
        # Split text into lines for processing
        lines = text.split('\n')
        
        current_participant = {}
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Try to extract name (look for lines with Hebrew characters)
            if re.search(r'[\u0590-\u05FF]', line) and not current_participant.get('name'):
                # Skip lines that contain activity types
                if not any(activity in line for activity in self.ACTIVITY_TYPES.keys()):
                    potential_name = re.search(hebrew_name_pattern, line)
                    if potential_name:
                        current_participant['name'] = potential_name.group().strip()
            
            # Extract ID number
            id_match = re.search(id_pattern, line)
            if id_match and not current_participant.get('id'):
                current_participant['id'] = id_match.group()
            
            # Extract receipt number
            receipt_match = re.search(receipt_pattern, line, re.IGNORECASE)
            if receipt_match and not current_participant.get('receipt_number'):
                current_participant['receipt_number'] = receipt_match.group(1)
            
            # Extract amount
            amount_match = re.search(amount_pattern, line)
            if amount_match and not current_participant.get('amount'):
                amount = amount_match.group(1) or amount_match.group(2)
                current_participant['amount'] = amount
            
            # If we have collected all information for a participant, add to list
            if all(key in current_participant for key in ['name', 'id', 'receipt_number', 'amount']):
                participants.append(current_participant.copy())
                current_participant = {}
        
        # Add any remaining participant if they have at least name and one other field
        if current_participant.get('name') and len(current_participant) > 1:
            participants.append(current_participant)
        
        return participants
    
    def process_text(self, text: str) -> Dict:
        """
        Process the input text to extract all relevant information.
        
        Args:
            text: The input Hebrew text
            
        Returns:
            Dictionary containing processed information
        """
        # Identify activity type
        self.activity_type = self.identify_activity_type(text)
        
        # Extract participant information
        self.participants = self.extract_participant_info(text)
        
        return {
            'activity_type': self.activity_type,
            'participants': self.participants,
            'total_participants': len(self.participants)
        }
    
    def create_word_document(self, output_path: str = None) -> str:
        """
        Create a Word document with the extracted information.
        
        Args:
            output_path: Path where to save the document
            
        Returns:
            Path to the created document
        """
        if not self.participants:
            raise ValueError("No participant information found. Please process text first.")
        
        # Create a new document
        doc = Document()
        
        # Set document direction to RTL for Hebrew
        doc.sections[0].start_type = 1
        
        # Create title
        current_month = self.HEBREW_MONTHS[datetime.now().month]
        
        # Use the first participant's name for the title, or a generic title
        if self.participants:
            title_name = self.participants[0]['name']
        else:
            title_name = "משתתפים"
        
        title = f"{title_name} - {current_month}"
        
        # Add title to document
        title_paragraph = doc.add_paragraph()
        title_run = title_paragraph.add_run(title)
        title_run.font.size = Inches(0.2)
        title_run.bold = True
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add some spacing
        doc.add_paragraph()
        
        # Create table
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        
        # Set table direction to RTL
        table.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Add header row
        header_cells = table.rows[0].cells
        headers = ['שם', 'תעודת זהות', 'מספר קבלה', 'סכום ששולם']
        
        for i, header in enumerate(headers):
            header_cells[i].text = header
            # Make header bold
            for paragraph in header_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add participant data
        for participant in self.participants:
            row_cells = table.add_row().cells
            row_cells[0].text = participant.get('name', '')
            row_cells[1].text = participant.get('id', '')
            row_cells[2].text = participant.get('receipt_number', '')
            row_cells[3].text = participant.get('amount', '')
            
            # Center align the data
            for cell in row_cells:
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Generate output filename if not provided
        if not output_path:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            activity_suffix = "post_birth" if "לאחר לידה" in str(self.activity_type) else "pregnancy"
            output_path = f"physiotherapy_form_{activity_suffix}_{timestamp}.docx"
        
        # Save the document
        doc.save(output_path)
        return output_path


def main():
    """Main function to run the application."""
    print("Hebrew Physiotherapy Workshops Form Generator")
    print("=" * 50)
    
    # Get input text
    print("\nPlease enter the Hebrew text (press Enter twice to finish):")
    lines = []
    while True:
        try:
            line = input()
            if line == "" and lines and lines[-1] == "":
                break
            lines.append(line)
        except EOFError:
            break
    
    input_text = "\n".join(lines)
    
    if not input_text.strip():
        print("No input text provided. Exiting.")
        return
    
    # Process the text
    processor = HebrewFormProcessor()
    result = processor.process_text(input_text)
    
    # Display results
    print(f"\nProcessing Results:")
    print(f"Activity Type: {result['activity_type']}")
    print(f"Total Participants: {result['total_participants']}")
    
    if result['participants']:
        print("\nExtracted Participant Information:")
        for i, participant in enumerate(result['participants'], 1):
            print(f"\nParticipant {i}:")
            for key, value in participant.items():
                print(f"  {key}: {value}")
    
    # Generate Word document
    try:
        output_file = processor.create_word_document()
        print(f"\nWord document created successfully: {output_file}")
    except Exception as e:
        print(f"Error creating Word document: {e}")


if __name__ == "__main__":
    main() 